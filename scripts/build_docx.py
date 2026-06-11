#!/usr/bin/env python3
"""把 Markdown 组装成中文学位论文风格 docx。

核心做法:
- 重定义 Word 内置 Heading1/2/3 样式(去主题字体、纯黑、保留大纲级别 outlineLvl),
  标题段落套这些样式 -> 进大纲/导航、可一键生成目录、且不变蓝(原内置是蓝色 accent)。
- 同时修 Hyperlink 样式为黑色无下划线 -> 生成的目录不会是蓝色下划线。
- 正文/表格/代码等用 run 级写死字体(eastAsia), 不依赖主题。
- 公式 $..$ / $$..$$ 转 OMML 可编辑公式, 字号字体写死。
- 所有格式参数来自 preset.py 的预设, 可由用户要求定制。
"""
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent))
from latex2omml import latex_to_omath  # noqa: E402
from md_parser import parse as parse_md  # noqa: E402
import preset as preset_mod  # noqa: E402

from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT  # noqa: E402
from docx.oxml import OxmlElement, parse_xml  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Cm, Pt  # noqa: E402

OMML_NS = ('xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math" '
           'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"')

ALIGN = {"center": WD_PARAGRAPH_ALIGNMENT.CENTER,
         "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
         "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
         "justify": WD_PARAGRAPH_ALIGNMENT.JUSTIFY}


def set_font(run, chinese="宋体", western="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = western
    run._element.rPr.rFonts.set(qn("w:eastAsia"), chinese)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_para(p, first_indent=False, align=None, before=0, after=0, line=1.25,
             indent_pt=24):
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.left_indent = Cm(0)
    pf.right_indent = Cm(0)
    pf.first_line_indent = Pt(indent_pt) if first_indent else Pt(0)
    if align is not None:
        p.alignment = align


# ===== 重定义内置标题/超链接样式 =====
def _rpr_clear_theme(rpr):
    """删除 rPr 里的主题字体引用与主题色, 改为显式写死。"""
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is not None:
        for attr in ("asciiTheme", "hAnsiTheme", "eastAsiaTheme", "cstheme"):
            a = qn(f"w:{attr}")
            if a in rFonts.attrib:
                del rFonts.attrib[a]


def _set_style_run(style, chinese, western, size, bold, color):
    rpr = style.element.get_or_add_rPr()
    _rpr_clear_theme(rpr)
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), western)
    rFonts.set(qn("w:hAnsi"), western)
    rFonts.set(qn("w:cs"), western)
    rFonts.set(qn("w:eastAsia"), chinese)
    # 字号
    for tag in ("w:sz", "w:szCs"):
        el = rpr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rpr.append(el)
        el.set(qn("w:val"), str(int(round(size * 2))))
    # 加粗
    for tag in ("w:b", "w:bCs"):
        el = rpr.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            rpr.append(el)
        el.set(qn("w:val"), "1" if bold else "0")
    # 颜色(纯黑, 去主题色)
    col = rpr.find(qn("w:color"))
    if col is None:
        col = OxmlElement("w:color")
        rpr.append(col)
    col.set(qn("w:val"), color)
    for a in ("themeColor", "themeShade", "themeTint"):
        if qn(f"w:{a}") in col.attrib:
            del col.attrib[qn(f"w:{a}")]


def _set_style_ppr(style, align, before, after, line, first_indent):
    ppr = style.element.get_or_add_pPr()
    sp = ppr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing")
        ppr.append(sp)
    sp.set(qn("w:before"), str(int(before * 20)))
    sp.set(qn("w:after"), str(int(after * 20)))
    sp.set(qn("w:line"), str(int(line * 240)))
    sp.set(qn("w:lineRule"), "auto")
    ind = ppr.find(qn("w:ind"))
    if first_indent:
        if ind is None:
            ind = OxmlElement("w:ind")
            ppr.append(ind)
        ind.set(qn("w:firstLineChars"), "200")
        ind.set(qn("w:firstLine"), "480")
    elif ind is not None:
        for a in ("firstLine", "firstLineChars"):
            if qn(f"w:{a}") in ind.attrib:
                del ind.attrib[qn(f"w:{a}")]
    jc = ppr.find(qn("w:jc"))
    if align:
        if jc is None:
            jc = OxmlElement("w:jc")
            ppr.append(jc)
        jc.set(qn("w:val"), align)


def redefine_heading_styles(doc, ps):
    """重定义内置 Heading1/2/3, 保留其 outlineLvl(进大纲), 去主题/蓝色。"""
    for lvl, h in ps["headings"].items():
        try:
            style = doc.styles[f"Heading {lvl}"]
        except KeyError:
            continue
        _set_style_run(style, h["chinese"], h["western"], h["size"],
                       h.get("bold", False), h.get("color", "000000"))
        _set_style_ppr(style, {"center": "center", "left": "left",
                               "right": "right"}.get(h.get("align", "left")),
                       h.get("before", 6), h.get("after", 6),
                       ps["body"]["line"], h.get("first_indent", False))


def fix_hyperlink_style(doc, ps):
    """把 Hyperlink 字符样式改成黑色无下划线, 这样生成目录不会变蓝带下划线。"""
    hl = ps.get("hyperlink", {})
    color = hl.get("color", "000000")
    underline = hl.get("underline", False)
    styles_el = doc.styles.element
    # 找或建 Hyperlink 样式
    target = None
    for st in styles_el.findall(qn("w:style")):
        if st.get(qn("w:styleId")) == "Hyperlink":
            target = st
            break
    if target is None:
        target = parse_xml(
            f'<w:style {OMML_NS} w:type="character" w:styleId="Hyperlink">'
            '<w:name w:val="Hyperlink"/><w:rPr></w:rPr></w:style>')
        styles_el.append(target)
    rpr = target.find(qn("w:rPr"))
    if rpr is None:
        rpr = OxmlElement("w:rPr")
        target.append(rpr)
    # 颜色
    col = rpr.find(qn("w:color"))
    if col is None:
        col = OxmlElement("w:color")
        rpr.append(col)
    col.set(qn("w:val"), color)
    for a in ("themeColor", "themeShade", "themeTint"):
        if qn(f"w:{a}") in col.attrib:
            del col.attrib[qn(f"w:{a}")]
    # 下划线
    u = rpr.find(qn("w:u"))
    if u is None:
        u = OxmlElement("w:u")
        rpr.append(u)
    u.set(qn("w:val"), "single" if underline else "none")


def setup_document(ps):
    doc = Document()
    pg = ps["page"]
    sec = doc.sections[0]
    sec.page_width = Cm(pg["width_cm"])
    sec.page_height = Cm(pg["height_cm"])
    sec.top_margin = Cm(pg["top_cm"])
    sec.bottom_margin = Cm(pg["bottom_cm"])
    sec.left_margin = Cm(pg["left_cm"])
    sec.right_margin = Cm(pg["right_cm"])
    b = ps["body"]
    style = doc.styles["Normal"]
    style.font.name = b["western"]
    style._element.rPr.rFonts.set(qn("w:eastAsia"), b["chinese"])
    style.font.size = Pt(b["size"])
    redefine_heading_styles(doc, ps)
    fix_hyperlink_style(doc, ps)
    return doc


# ===== 行内解析: 文本 / 公式 / 粗体 / 斜体 / 等宽 =====
def _split_inline(text):
    parts = []
    pos = 0
    for m in re.finditer(r"(?<!\\)\$([^$]+?)(?<!\\)\$", text):
        if m.start() > pos:
            parts.extend(_split_emphasis(text[pos:m.start()]))
        parts.append(("math", m.group(1).strip()))
        pos = m.end()
    if pos < len(text):
        parts.extend(_split_emphasis(text[pos:]))
    return parts


def _split_emphasis(text):
    parts = []
    pos = 0
    pat = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`([^`]+?)`")
    for m in pat.finditer(text):
        if m.start() > pos:
            parts.append(("text", text[pos:m.start()]))
        if m.group(1) is not None:
            parts.append(("bold", m.group(1)))
        elif m.group(2) is not None:
            parts.append(("italic", m.group(2)))
        else:
            parts.append(("mono", m.group(3)))
        pos = m.end()
    if pos < len(text):
        parts.append(("text", text[pos:]))
    return parts


def _add_omath(p, latex, ps):
    omath = latex_to_omath(latex, size_pt=ps["math"]["size"])
    inner = omath[len("<m:oMath>"):-len("</m:oMath>")]
    p._p.append(parse_xml(f'<m:oMath {OMML_NS}>{inner}</m:oMath>'))


def _render_inline(p, text, ps, size=None):
    b = ps["body"]
    size = size or b["size"]
    for kind, content in _split_inline(text):
        if kind == "math":
            _add_omath(p, content, ps)
        elif kind == "bold":
            set_font(p.add_run(content), b["chinese"], b["western"], size, bold=True)
        elif kind == "italic":
            set_font(p.add_run(content), b["chinese"], b["western"], size, italic=True)
        elif kind == "mono":
            set_font(p.add_run(content), b["chinese"], ps["code"]["western"], size)
        else:
            set_font(p.add_run(content), b["chinese"], b["western"], size)


# ===== 块渲染 =====
def add_heading(doc, level, text, ps):
    lvl = min(level, 3)
    p = doc.add_paragraph(style=f"Heading {lvl}")  # 套内置样式 -> 进大纲
    # 样式已定字体/字号/对齐/间距; 这里只补正文行距, run 不再写死(交给样式)
    set_para(p, first_indent=ps["headings"][lvl].get("first_indent", False),
             align=ALIGN.get(ps["headings"][lvl].get("align", "left")),
             before=ps["headings"][lvl].get("before", 6),
             after=ps["headings"][lvl].get("after", 6),
             line=ps["body"]["line"])
    p.add_run(text)


def add_paragraph(doc, text, ps):
    p = doc.add_paragraph()
    set_para(p, first_indent=True, line=ps["body"]["line"],
             indent_pt=int(ps["body"]["size"] * ps["body"].get("first_indent_chars", 2)))
    _render_inline(p, text, ps)


def add_mathblock(doc, latex, ps):
    p = doc.add_paragraph()
    set_para(p, first_indent=False, align=WD_PARAGRAPH_ALIGNMENT.CENTER,
             before=6, after=6, line=ps["body"]["line"])
    _add_omath(p, latex, ps)


def add_list(doc, items, ordered, ps):
    lst = ps["list"]
    for idx, item in enumerate(items, 1):
        p = doc.add_paragraph()
        set_para(p, first_indent=False, line=lst["line"])
        p.paragraph_format.left_indent = Cm(lst["left_indent_cm"])
        p.paragraph_format.first_line_indent = Cm(-lst.get("hanging_cm", 0.37))
        prefix = f"{idx}. " if ordered else "· "
        set_font(p.add_run(prefix), lst["chinese"], lst["western"], lst["size"])
        _render_inline(p, item, ps, size=lst["size"])


def add_quote(doc, text, ps):
    q = ps["quote"]
    p = doc.add_paragraph()
    set_para(p, first_indent=False, line=q["line"])
    p.paragraph_format.left_indent = Cm(q["left_indent_cm"])
    _render_inline(p, text, ps, size=q["size"])


def add_code(doc, text, ps):
    c = ps["code"]
    for line in text.rstrip("\n").split("\n"):
        p = doc.add_paragraph()
        set_para(p, first_indent=False, line=c["line"])
        p.paragraph_format.left_indent = Cm(c["left_indent_cm"])
        set_font(p.add_run(line if line else " "), c["chinese"], c["western"], c["size"])


def _cell_border(cell, top=None, bottom=None):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "bottom", "left", "right", "insideH", "insideV"):
        el = tcBorders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tcBorders.append(el)
        if edge == "top" and top:
            for k, v in top.items():
                el.set(qn(f"w:{k}"), str(v))
        elif edge == "bottom" and bottom:
            for k, v in bottom.items():
                el.set(qn(f"w:{k}"), str(v))
        else:
            el.set(qn("w:val"), "nil")


def add_table(doc, rows, ps):
    tb = ps["table"]
    n_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    color = tb["line_color"]
    thick = {"val": "single", "sz": tb["top_sz"], "space": "0", "color": color}
    head = {"val": "single", "sz": tb["header_sz"], "space": "0", "color": color}
    bot = {"val": "single", "sz": tb["bottom_sz"], "space": "0", "color": color}
    last = len(rows) - 1
    for ri, row in enumerate(rows):
        for ci in range(n_cols):
            val = row[ci] if ci < len(row) else ""
            cell = table.cell(ri, ci)
            cell.text = ""
            cp = cell.paragraphs[0]
            set_para(cp, first_indent=False, align=WD_PARAGRAPH_ALIGNMENT.CENTER,
                     line=tb["cell_line"])
            set_font(cp.add_run(val), tb["chinese"], tb["western"], tb["size"],
                     bold=(ri == 0 and tb["header_bold"]))
            top = thick if ri == 0 else None
            bottom = head if ri == 0 else (bot if ri == last else None)
            _cell_border(cell, top=top, bottom=bottom)
    doc.add_paragraph()


def add_image(doc, path, alt, ps):
    p = doc.add_paragraph()
    set_para(p, first_indent=False, align=WD_PARAGRAPH_ALIGNMENT.CENTER, before=6, after=3)
    try:
        p.add_run().add_picture(path, width=Cm(12.0))
    except Exception:
        set_font(p.add_run(f"[图片缺失: {path}]"), size=10.5)
    if alt:
        cap = doc.add_paragraph()
        set_para(cap, first_indent=False, align=WD_PARAGRAPH_ALIGNMENT.CENTER, after=6)
        cp = ps["caption"]
        set_font(cap.add_run(alt), cp["chinese"], cp["western"], cp["size"])


# ===== 主流程 =====
def build(md_path, out_path, title=None, author=None, ps=None):
    if ps is None:
        ps = preset_mod.get_default()
    md = Path(md_path).read_text(encoding="utf-8")
    blocks = parse_md(md)
    doc = setup_document(ps)

    if title:
        t = ps["title"]
        p = doc.add_paragraph()
        set_para(p, first_indent=False, align=ALIGN.get(t.get("align", "center")), after=6)
        set_font(p.add_run(title), t["chinese"], t["western"], t["size"], bold=t.get("bold", False))
    if author:
        a = ps["author"]
        p = doc.add_paragraph()
        set_para(p, first_indent=False, align=ALIGN.get(a.get("align", "center")), after=12)
        set_font(p.add_run(author), a["chinese"], a["western"], a["size"])

    for blk in blocks:
        t = blk["t"]
        if t == "h":
            add_heading(doc, blk["level"], blk["text"], ps)
        elif t == "p":
            add_paragraph(doc, blk["text"], ps)
        elif t == "mathblock":
            add_mathblock(doc, blk["latex"], ps)
        elif t == "list":
            add_list(doc, blk["items"], blk["ordered"], ps)
        elif t == "quote":
            add_quote(doc, blk["text"], ps)
        elif t == "code":
            add_code(doc, blk["text"], ps)
        elif t == "table":
            add_table(doc, blk["rows"], ps)
        elif t == "image":
            add_image(doc, blk["path"], blk["alt"], ps)

    doc.save(out_path)
    # 字体存在性检查: 缺失则警告(不阻断), 避免 Word 静默替换字体却无人察觉
    try:
        import fontcheck
        fontcheck.warn_if_missing(ps)
    except Exception:
        pass
    return out_path


if __name__ == "__main__":
    import argparse
    ap = argparse.ArgumentParser()
    ap.add_argument("input")
    ap.add_argument("-o", "--output")
    ap.add_argument("--title")
    ap.add_argument("--author")
    ap.add_argument("--preset", help="自定义 preset JSON 路径")
    a = ap.parse_args()
    out = a.output or str(Path(a.input).with_suffix(".docx"))
    ps = preset_mod.load(a.preset) if a.preset else preset_mod.get_default()
    build(a.input, out, title=a.title, author=a.author, ps=ps)
    print(f"完成 -> {out}")

